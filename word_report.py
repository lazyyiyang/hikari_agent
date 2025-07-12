#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
from datetime import datetime, timedelta
import io

def create_research_report():
    """创建研报样式的Word文档"""
    
    # 创建新文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # 添加页眉信息
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run('证券研究报告 | 首次覆盖报告')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 添加日期
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_p.add_run(f'{datetime.now().strftime("%Y 年 %m 月 %d 日")}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run('智能科技 (08888.HK)')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 153)
    
    # 添加副标题
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_p.add_run('AI驱动未来，构建智能生态平台')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(51, 51, 51)
    
    # 添加空行
    doc.add_paragraph()
    
    # 创建主要内容和信息表格的表格
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.columns[0].width = Inches(4.5)
    main_table.columns[1].width = Inches(2.0)
    
    # 左侧主要内容
    left_cell = main_table.cell(0, 0)
    left_p = left_cell.paragraphs[0]
    left_p.text = ""
    
    # 添加主要内容
    content_text = """作为新兴的人工智能和大数据技术公司，智能科技专注于机器学习、深度学习、自然语言处理等前沿技术领域，致力于为企业客户提供智能化解决方案。

2024年公司营收达到38.5亿元，同比增长65%，其中AI产品收入占比达到78%。公司在智能语音、计算机视觉、自然语言处理等核心技术领域均有重大突破。

随着数字化转型加速和AI技术的快速发展，我们预计公司将在未来三年内实现快速增长。预计2030年市场规模将达到2000亿元，公司有望占据10%的市场份额。

在技术创新方面，公司已获得156项技术专利，在智能语音识别准确率达到98.5%，计算机视觉识别精度达到99.2%，均处于行业领先水平。

基于公司强大的技术实力和广阔的市场前景，我们维持对公司的"买入"评级，目标价格为12.5港元。"""
    
    left_p.add_run(content_text)
    left_p.style.font.size = Pt(10)
    
    # 右侧信息表格
    right_cell = main_table.cell(0, 1)
    right_p = right_cell.paragraphs[0]
    right_p.text = ""
    
    # 创建信息表格
    info_table = right_cell.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    
    # 添加表头
    header_cell = info_table.cell(0, 0)
    header_cell.merge(info_table.cell(0, 1))
    header_p = header_cell.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run('买入（首次）')
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    # 设置背景色为蓝色
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '4472C4')
    header_cell._tc.get_or_add_tcPr().append(shading)
    
    # 添加信息数据
    info_data = [
        ['', ''],  # 已用作表头
        ['股票信息', ''],
        ['行业', '科技'],
        ['06月30日收盘价（港元）', '8.45'],
        ['总市值（百万港元）', '25,680.33'],
        ['总股本（百万股）', '3,040.25'],
        ['流通市值占比（%）', '85.6'],
        ['30日日均成交额（百万港元）', '125.8']
    ]
    
    for i, (key, value) in enumerate(info_data[1:], 1):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value
        
        # 设置字体大小
        for j in range(2):
            for paragraph in info_table.cell(i, j).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加图表标题
    chart_title = doc.add_paragraph()
    chart_title.add_run('股价走势').font.bold = True
    chart_title.add_run('').font.size = Pt(12)
    
    # 创建股价走势图
    create_stock_chart()
    
    # 插入图表（这里用文字代替，实际使用时需要插入图片）
    chart_p = doc.add_paragraph()
    chart_p.add_run('[股价走势图占位符 - 实际使用时请插入生成的图表]')
    chart_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加分析师信息
    doc.add_paragraph()
    analyst_title = doc.add_paragraph()
    analyst_title.add_run('分析师').font.bold = True
    
    analyst_info = doc.add_table(rows=3, cols=2)
    analyst_data = [
        ['首席分析师', '张明'],
        ['执业证书编号', 'S0680519100005'],
        ['邮箱', 'zhangming@example.com']
    ]
    
    for i, (key, value) in enumerate(analyst_data):
        analyst_info.cell(i, 0).text = key
        analyst_info.cell(i, 1).text = value
    
    # 添加财务预测表
    doc.add_paragraph()
    forecast_title = doc.add_paragraph()
    forecast_title.add_run('财务预测').font.bold = True
    
    forecast_table = doc.add_table(rows=9, cols=6)
    forecast_table.style = 'Table Grid'
    
    # 财务预测数据
    forecast_data = [
        ['财务指标', '2023A', '2024A', '2025E', '2026E', '2027E'],
        ['营业收入（百万元）', '2,340', '3,850', '5,200', '7,800', '11,200'],
        ['增长率YoY（%）', '45', '65', '35', '50', '44'],
        ['净利润（百万元）', '234', '462', '780', '1,248', '1,904'],
        ['净利润增长率（%）', '78', '97', '69', '60', '53'],
        ['毛利率（%）', '45.2', '48.5', '52.0', '54.5', '56.0'],
        ['EPS（港元）', '0.08', '0.15', '0.26', '0.41', '0.63'],
        ['P/E（倍）', '106', '56', '33', '21', '13'],
        ['P/S（倍）', '11.0', '6.7', '4.9', '3.3', '2.3']
    ]
    
    for i, row_data in enumerate(forecast_data):
        for j, cell_data in enumerate(row_data):
            cell = forecast_table.cell(i, j)
            cell.text = cell_data
            
            # 设置表头样式
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                # 设置表头背景色
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E7E6E6')
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    # 添加风险提示
    doc.add_paragraph()
    risk_title = doc.add_paragraph()
    risk_title.add_run('风险提示：').font.bold = True
    
    risk_content = doc.add_paragraph()
    risk_content.add_run('技术发展不及预期风险、市场竞争加剧风险、政策变化风险、汇率波动风险等。请投资者谨慎决策。')
    
    # 添加页脚
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('请仔细阅读本报告末页重要声明')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    doc.save('研报样式报告.docx')
    print("研报已生成：研报样式报告.docx")

def create_stock_chart():
    """创建股价走势图"""
    # 设置中文字体
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    
    # 生成模拟数据
    dates = [datetime.now() - timedelta(days=x) for x in range(180, 0, -1)]
    
    # 生成股价数据
    np.random.seed(42)
    base_price = 6.5
    returns = np.random.normal(0.001, 0.03, 180)
    prices = [base_price]
    
    for ret in returns:
        prices.append(prices[-1] * (1 + ret))
    
    # 生成指数数据
    index_returns = np.random.normal(0.0005, 0.02, 180)
    index_prices = [100]
    
    for ret in index_returns:
        index_prices.append(index_prices[-1] * (1 + ret))
    
    # 创建图表
    fig, ax = plt.subplots(figsize=(8, 4))
    
    # 绘制股价和指数
    ax.plot(dates, prices[1:], label='智能科技', linewidth=2, color='blue')
    ax.plot(dates, [p/index_prices[1]*prices[1] for p in index_prices[1:]], 
            label='恒生指数', linewidth=1, color='gray', linestyle='--')
    
    # 设置图表样式
    ax.set_xlabel('日期')
    ax.set_ylabel('价格/指数')
    ax.legend()
    ax.grid(True, alpha=0.3)
    
    # 设置日期格式
    import matplotlib.dates as mdates
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m'))
    ax.xaxis.set_major_locator(mdates.MonthLocator(interval=1))
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=45)
    
    plt.title('股价走势对比图')
    plt.tight_layout()
    plt.savefig('stock_chart.png', dpi=300, bbox_inches='tight')
    plt.close()
    
    print("股价走势图已生成：stock_chart.png")

def main():
    """主函数"""
    print("开始生成研报样式Word文档...")
    
    try:
        create_research_report()
        print("研报生成完成！")
        print("\n生成的文件：")
        print("1. 研报样式报告.docx - 主要报告文档")
        print("2. stock_chart.png - 股价走势图")
        print("\n注意：")
        print("- 请确保已安装 python-docx 和 matplotlib 库")
        print("- 如需插入图表到Word中，请手动插入生成的PNG图片")
        print("- 可以根据需要调整字体、颜色和格式")
        
    except Exception as e:
        print(f"生成过程中出现错误：{e}")
        print("请确保已安装所需的依赖库：")
        print("pip install python-docx matplotlib numpy")

if __name__ == "__main__":
    main()